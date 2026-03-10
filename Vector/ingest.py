"""
ingest.py（複数ファイル対応版）
================================
指定フォルダ内のPDF・Word・TXTを一括読み込み → チャンク分割 → Gemini Embeddingでベクトル化 → FAISSに保存

実行方法:
    # フォルダごと一括取り込み（推奨）
    python ingest.py --folder ./docs

    # 単一ファイル指定
    python ingest.py --pdf 就業規則.pdf

    # 複数ファイル個別指定
    python ingest.py --pdf 規定A.pdf 規定B.pdf マニュアル.pdf

対応ファイル形式:
    .pdf  / .docx / .txt / .md
"""

import os
import argparse
import glob
from pathlib import Path

from pypdf import PdfReader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain_community.vectorstores import FAISS
from langchain.schema import Document

# ===================== 設定 =====================

FAISS_DIR     = "./faiss_index"
EMBED_MODEL   = "models/gemini-embedding-001"
CHUNK_SIZE    = 500
CHUNK_OVERLAP = 100

# 対応拡張子
SUPPORTED_EXTENSIONS = [".pdf", ".docx", ".txt", ".md"]


# ===================== STEP 1: ファイル収集 =====================

def collect_files(folder: str = None, pdf_files: list = None) -> list[Path]:
    """
    フォルダまたは個別指定からファイルパスのリストを返す。

    Args:
        folder:    フォルダパス（指定時はフォルダ内を再帰的に検索）
        pdf_files: 個別ファイルパスのリスト

    Returns:
        Pathオブジェクトのリスト
    """
    paths = []

    if folder:
        print(f"\n📁 フォルダを検索中: {folder}")
        for ext in SUPPORTED_EXTENSIONS:
            # ** で再帰的にサブフォルダも検索
            found = glob.glob(f"{folder}/**/*{ext}", recursive=True)
            paths.extend([Path(p) for p in found])

        if not paths:
            print(f"   ⚠️  対応ファイルが見つかりません")
            print(f"   対応形式: {', '.join(SUPPORTED_EXTENSIONS)}")
            return []

        print(f"   ✅ {len(paths)}個のファイルを発見")
        for p in paths:
            print(f"      - {p.name}  ({p.suffix})")

    elif pdf_files:
        paths = [Path(p) for p in pdf_files]

    return paths


# ===================== STEP 2: テキスト抽出（形式別） =====================

def extract_text_from_pdf(path: Path) -> str:
    """PDFからテキスト抽出"""
    reader = PdfReader(str(path))
    text = ""
    for i, page in enumerate(reader.pages):
        page_text = page.extract_text()
        if page_text:
            text += f"\n[ページ {i+1}]\n{page_text}"
    return text


def extract_text_from_docx(path: Path) -> str:
    """Wordファイルからテキスト抽出"""
    try:
        import docx
        doc = docx.Document(str(path))
        return "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
    except ImportError:
        print(f"   ⚠️  python-docx が未インストール: pip install python-docx")
        return ""


def extract_text_from_txt(path: Path) -> str:
    """テキスト・Markdownファイルからテキスト抽出"""
    encodings = ["utf-8", "utf-8-sig", "cp932", "shift_jis"]
    for enc in encodings:
        try:
            return path.read_text(encoding=enc)
        except (UnicodeDecodeError, LookupError):
            continue
    print(f"   ⚠️  文字コードを判定できませんでした: {path.name}")
    return ""


def extract_text(path: Path) -> str:
    """ファイル形式に応じてテキスト抽出を振り分ける"""
    ext = path.suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(path)
    elif ext == ".docx":
        return extract_text_from_docx(path)
    elif ext in [".txt", ".md"]:
        return extract_text_from_txt(path)
    else:
        print(f"   ⚠️  未対応形式: {path.name}")
        return ""


# ===================== STEP 3: 全ファイルを読み込んでDocumentリストを生成 =====================

def load_all_documents(paths: list[Path]) -> list[Document]:
    """
    全ファイルを読み込み、LangChainのDocumentオブジェクトに変換する。

    Document には以下のメタデータを付与：
        - source:    ファイル名（回答時の出典表示に使う）
        - file_type: 拡張子
        - file_path: フルパス

    Returns:
        Documentオブジェクトのリスト
    """
    print(f"\n📄 ファイルを読み込み中...")
    documents = []
    total_chars = 0

    for path in paths:
        print(f"   読み込み: {path.name}", end="")
        text = extract_text(path)

        if not text.strip():
            print(f"  → スキップ（テキストなし）")
            continue

        # ファイル名をソースとして記録
        # → 後で「このチャンクはどのファイルから来たか」がわかる
        doc = Document(
            page_content=text,
            metadata={
                "source":    path.name,
                "file_type": path.suffix.lower(),
                "file_path": str(path),
            }
        )
        documents.append(doc)
        total_chars += len(text)
        print(f"  → {len(text):,}文字")

    print(f"\n   ✅ {len(documents)}ファイル読み込み完了 / 総文字数: {total_chars:,}文字")
    return documents


# ===================== STEP 4: チャンク分割 =====================

def split_documents(documents: list[Document]) -> list[Document]:
    """
    各Documentをチャンクに分割する。

    メタデータ（source等）は各チャンクに引き継がれる。
    → どのファイルのどの部分から来たチャンクかが追跡できる。

    Returns:
        チャンク化されたDocumentのリスト
    """
    print(f"\n✂️  チャンクに分割中...")

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        separators=["\n\n", "\n", "。", "、", " ", ""],
    )

    chunks = splitter.split_documents(documents)

    print(f"   ✅ {len(chunks)}チャンクに分割完了")
    print(f"   📊 チャンクサイズ: {CHUNK_SIZE}文字 / オーバーラップ: {CHUNK_OVERLAP}文字")

    # ファイル別チャンク数を表示
    source_counts = {}
    for chunk in chunks:
        src = chunk.metadata.get("source", "不明")
        source_counts[src] = source_counts.get(src, 0) + 1

    print(f"\n   【ファイル別チャンク数】")
    for src, count in source_counts.items():
        print(f"      {src}: {count}チャンク")

    return chunks


# ===================== STEP 5: ベクトル化 & FAISS保存 =====================

def embed_and_store(chunks: list[Document], api_key: str) -> FAISS:
    """
    チャンクをGemini Embeddingでベクトル化してFAISSに保存する。

    既存のFAISSインデックスがある場合は追記モードで保存。
    → 後からドキュメントを追加していける。
    """
    print(f"\n🔢 Gemini Embeddingでベクトル化中...")
    print(f"   使用モデル: {EMBED_MODEL}")
    print(f"   チャンク数: {len(chunks)}個")

    embeddings = GoogleGenerativeAIEmbeddings(
        model=EMBED_MODEL,
        google_api_key=api_key,
    )

    # テキストとメタデータを分離
    texts     = [c.page_content for c in chunks]
    metadatas = [c.metadata      for c in chunks]

    existing_index = Path(FAISS_DIR)

    if existing_index.exists():
        # ===================== 追記モード =====================
        # 既存インデックスに新しいチャンクを追加する
        # → 「今あるナレッジに追加」ができる
        print(f"\n💾 既存FAISSインデックスに追記中...")
        vectorstore = FAISS.load_local(
            FAISS_DIR,
            embeddings,
            allow_dangerous_deserialization=True,
        )
        vectorstore.add_texts(texts=texts, metadatas=metadatas)
        vectorstore.save_local(FAISS_DIR)
        print(f"   ✅ 追記完了！")

    else:
        # ===================== 新規作成モード =====================
        print(f"\n💾 FAISSインデックスを新規作成中...")
        vectorstore = FAISS.from_texts(
            texts=texts,
            embedding=embeddings,
            metadatas=metadatas,
        )
        vectorstore.save_local(FAISS_DIR)
        print(f"   ✅ 新規作成完了！")

    print(f"   📁 保存場所: {Path(FAISS_DIR).resolve()}")
    return vectorstore


# ===================== STEP 6: 保存確認 =====================

def verify_storage(vectorstore: FAISS):
    """試験的な類似検索で動作確認"""
    print(f"\n🔍 動作確認: 試験的な類似検索を実行...")

    test_queries = ["この文書の主なトピックは何ですか？", "重要なルールや規則を教えてください"]

    for query in test_queries:
        results = vectorstore.similarity_search_with_score(query, k=2)
        if results:
            doc, score = results[0]
            source = doc.metadata.get("source", "不明")
            preview = doc.page_content[:50].replace('\n', ' ')
            print(f"\n   Q: 「{query}」")
            print(f"   → ソース: {source} / スコア: {score:.4f}")
            print(f"   → チャンク: 「{preview}...」")


# ===================== メイン処理 =====================

def main():
    parser = argparse.ArgumentParser(
        description="複数ファイルをRAG用にベクトル化してFAISSに保存",
        formatter_class=argparse.RawTextHelpFormatter,
    )
    parser.add_argument(
        "--folder", "-f",
        help="フォルダパス（フォルダ内の対応ファイルを一括取り込み）\n例: --folder ./docs"
    )
    parser.add_argument(
        "--pdf", "-p",
        nargs="+",
        help="ファイルパス（複数指定可）\n例: --pdf 規定A.pdf 規定B.pdf マニュアル.docx"
    )
    parser.add_argument(
        "--reset",
        action="store_true",
        help="既存のFAISSインデックスを削除してから新規作成"
    )
    args = parser.parse_args()

    if not args.folder and not args.pdf:
        parser.print_help()
        return

    # 既存インデックスのリセット
    if args.reset:
        import shutil
        if Path(FAISS_DIR).exists():
            shutil.rmtree(FAISS_DIR)
            print(f"🗑️  既存インデックスを削除しました")

    # APIキー取得
    api_key = os.environ.get("GEMINI_API_KEY")
    if not api_key:
        api_key = input("Gemini APIキーを入力してください: ").strip()

    print("=" * 60)
    print("🚀 RAG Ingest（複数ファイル対応版）開始")
    print("=" * 60)

    # パイプライン実行
    paths     = collect_files(folder=args.folder, pdf_files=args.pdf)
    if not paths:
        return

    documents = load_all_documents(paths)
    if not documents:
        print("❌ 読み込めるファイルがありませんでした")
        return

    chunks    = split_documents(documents)
    db        = embed_and_store(chunks, api_key)
    verify_storage(db)

    print("\n" + "=" * 60)
    print("✅ ingest完了！")
    print("   streamlit run app.py でアプリを起動してください")
    print("=" * 60)


if __name__ == "__main__":
    main()